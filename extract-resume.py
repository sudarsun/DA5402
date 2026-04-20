from __future__ import annotations
import argparse
import json
import re
from collections import OrderedDict
from datetime import datetime
from pathlib import Path
from typing import Any
import pdfplumber
from docx import Document

#!/usr/bin/env python3
"""
Resume parser for PDF, DOC, and DOCX files.

Purpose:
- Extract common resume fields at scale
- Intentionally excludes protected/sensitive demographic attributes
    such as age, gender, race, religion, marital status, disability, etc.

Install:
        pip install pdfplumber python-docx
Optional for .doc support:
        pip install textract

Usage:
        python extract_resumes.py /path/to/resumes -o extracted_resumes.json
"""




SUPPORTED_EXTENSIONS = {".pdf", ".doc", ".docx"}

MONTHS = {
        "jan": 1,
        "january": 1,
        "feb": 2,
        "february": 2,
        "mar": 3,
        "march": 3,
        "apr": 4,
        "april": 4,
        "may": 5,
        "jun": 6,
        "june": 6,
        "jul": 7,
        "july": 7,
        "aug": 8,
        "august": 8,
        "sep": 9,
        "sept": 9,
        "september": 9,
        "oct": 10,
        "october": 10,
        "nov": 11,
        "november": 11,
        "dec": 12,
        "december": 12,
}

COMMON_SKILLS = {
        "python", "java", "javascript", "typescript", "c", "c++", "c#", "go", "rust", "scala",
        "ruby", "php", "swift", "kotlin", "r", "sql", "mysql", "postgresql", "sqlite", "oracle",
        "mongodb", "redis", "snowflake", "power bi", "tableau", "excel", "pandas", "numpy",
        "scikit-learn", "tensorflow", "pytorch", "machine learning", "deep learning", "nlp",
        "data analysis", "data science", "data engineering", "etl", "spark", "hadoop", "airflow",
        "docker", "kubernetes", "linux", "git", "github", "azure", "aws", "gcp", "terraform",
        "ansible", "jenkins", "ci/cd", "devops", "react", "angular", "vue", "node.js", "django",
        "flask", "fastapi", "spring", "spring boot", "rest", "graphql", "html", "css", "sass",
        "tailwind", "bootstrap", "microservices", "agile", "scrum", "jira", "selenium",
        "pytest", "unit testing", "api testing", "cybersecurity", "networking", "salesforce",
        "sap", "project management", "product management", "business analysis", "figma",
        "ui/ux", "leadership", "communication"
}

COMMON_LANGUAGES = {
        "english", "spanish", "french", "german", "italian", "portuguese", "hindi", "arabic",
        "mandarin", "chinese", "japanese", "korean", "russian", "urdu", "bengali", "tamil",
        "telugu", "marathi", "punjabi", "gujarati", "dutch", "turkish", "polish"
}

DEGREE_PATTERN = re.compile(
        r"\b("
        r"b\.?s\.?|bachelor(?:'s)?|b\.?a\.?|master(?:'s)?|m\.?s\.?|m\.?a\.?|mba|"
        r"ph\.?d\.?|doctorate|associate(?:'s)?|btech|mtech|be|me|bca|mca"
        r")\b",
        re.IGNORECASE,
)

INSTITUTION_PATTERN = re.compile(
        r"\b(university|college|institute|school|academy|polytechnic)\b",
        re.IGNORECASE,
)

EMAIL_PATTERN = re.compile(r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b", re.IGNORECASE)
URL_PATTERN = re.compile(r"(https?://[^\s]+|www\.[^\s]+)", re.IGNORECASE)

PHONE_CANDIDATE_PATTERN = re.compile(
        r"(?<!\d)(?:\+?\d{1,3}[\s().-]*)?(?:\(?\d{2,4}\)?[\s().-]*){2,5}\d(?!\d)"
)

DATE_RANGE_PATTERN = re.compile(
        r"(?P<start>(?:"
        r"(?:jan|january|feb|february|mar|march|apr|april|may|jun|june|jul|july|aug|august|"
        r"sep|sept|september|oct|october|nov|november|dec|december)\s+\d{4}"
        r"|\d{4}))"
        r"\s*(?:-|–|—|to)\s*"
        r"(?P<end>(?:present|current|now|"
        r"(?:jan|january|feb|february|mar|march|apr|april|may|jun|june|jul|july|aug|august|"
        r"sep|sept|september|oct|october|nov|november|dec|december)\s+\d{4}"
        r"|\d{4}))",
        re.IGNORECASE,
)


def unique(seq: list[str]) -> list[str]:
        return list(OrderedDict.fromkeys(seq))


def normalize_text(text: str) -> str:
        text = text.replace("\x00", " ")
        text = text.replace("\uf0b7", " ")
        text = text.replace("\u2022", " ")
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()


def extract_text_from_pdf(path: Path) -> str:
        chunks = []
        with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                        page_text = page.extract_text() or ""
                        if page_text:
                                chunks.append(page_text)
        return normalize_text("\n".join(chunks))


def extract_text_from_docx(path: Path) -> str:
        doc = Document(path)
        chunks = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
                for row in table.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                        if row_text:
                                chunks.append(row_text)
        return normalize_text("\n".join(chunks))


def extract_text_from_doc(path: Path) -> str:
        try:
                import textract  # type: ignore

                raw = textract.process(str(path))
                return normalize_text(raw.decode("utf-8", errors="ignore"))
        except Exception as exc:
                raise RuntimeError(
                        "Unable to parse .doc file. Install 'textract' and required system dependencies."
                ) from exc


def extract_text(path: Path) -> str:
        suffix = path.suffix.lower()
        if suffix == ".pdf":
                return extract_text_from_pdf(path)
        if suffix == ".docx":
                return extract_text_from_docx(path)
        if suffix == ".doc":
                return extract_text_from_doc(path)
        raise ValueError(f"Unsupported file type: {path.suffix}")


def clean_lines(text: str) -> list[str]:
        lines = []
        for line in text.splitlines():
                line = re.sub(r"\s+", " ", line).strip(" |,-")
                if line:
                        lines.append(line)
        return lines


def extract_emails(text: str) -> list[str]:
        return unique([m.group(0).strip(".,;") for m in EMAIL_PATTERN.finditer(text)])


def extract_urls(text: str) -> list[str]:
        urls = []
        for match in URL_PATTERN.finditer(text):
                url = match.group(0).strip(".,;)")
                if url.lower().startswith("www."):
                        url = "https://" + url
                urls.append(url)
        return unique(urls)


def extract_phones(text: str) -> list[str]:
        phones = []
        for match in PHONE_CANDIDATE_PATTERN.finditer(text):
                raw = match.group(0).strip()
                digits = re.sub(r"\D", "", raw)
                if 10 <= len(digits) <= 15:
                        phones.append(raw)
        return unique(phones)


def extract_links(urls: list[str]) -> dict[str, Any]:
        links = {
                "linkedin": None,
                "github": None,
                "portfolio": [],
                "other": [],
        }
        for url in urls:
                lower = url.lower()
                if "linkedin.com" in lower and not links["linkedin"]:
                        links["linkedin"] = url
                elif "github.com" in lower and not links["github"]:
                        links["github"] = url
                elif any(x in lower for x in ["portfolio", "behance", "dribbble"]):
                        links["portfolio"].append(url)
                else:
                        links["other"].append(url)

        links["portfolio"] = unique(links["portfolio"])
        links["other"] = unique(links["other"])
        return links


def extract_name(lines: list[str]) -> str | None:
        stop_words = {
                "resume", "curriculum vitae", "cv", "email", "phone", "mobile", "address",
                "linkedin", "github", "profile", "summary", "objective"
        }

        for line in lines[:10]:
                candidate = line.strip()
                lower = candidate.lower()
                if any(word in lower for word in stop_words):
                        continue
                if "@" in candidate or re.search(r"\d", candidate):
                        continue
                if len(candidate) > 60 or len(candidate.split()) < 2 or len(candidate.split()) > 5:
                        continue

                tokens = re.findall(r"[A-Za-z][A-Za-z'`.-]+", candidate)
                if 2 <= len(tokens) <= 5:
                        return " ".join(token.strip(".,") for token in tokens)
        return None


def extract_location(text: str, lines: list[str]) -> str | None:
        patterns = [
                re.compile(r"\b[A-Z][a-z]+(?: [A-Z][a-z]+)*,\s*[A-Z]{2}\b"),
                re.compile(r"\b[A-Z][a-z]+(?: [A-Z][a-z]+)*,\s*[A-Z][a-z]+(?: [A-Z][a-z]+)*\b"),
        ]

        for line in lines[:20]:
                if "@" in line or "http" in line.lower():
                        continue
                for pattern in patterns:
                        m = pattern.search(line)
                        if m:
                                return m.group(0)

        for pattern in patterns:
                m = pattern.search(text)
                if m:
                        return m.group(0)

        return None


def extract_skills(text: str) -> list[str]:
        found = []
        haystack = f" {text.lower()} "
        for skill in sorted(COMMON_SKILLS, key=len, reverse=True):
                pattern = r"(?<!\w)" + re.escape(skill.lower()) + r"(?!\w)"
                if re.search(pattern, haystack, re.IGNORECASE):
                        found.append(skill)
        return unique(found)


def extract_languages(text: str) -> list[str]:
        found = []
        haystack = f" {text.lower()} "
        for language in COMMON_LANGUAGES:
                pattern = r"(?<!\w)" + re.escape(language.lower()) + r"(?!\w)"
                if re.search(pattern, haystack, re.IGNORECASE):
                        found.append(language.title())
        return sorted(unique(found))


def extract_education(lines: list[str]) -> list[dict[str, Any]]:
        education = []
        for i, line in enumerate(lines):
                if DEGREE_PATTERN.search(line) or INSTITUTION_PATTERN.search(line):
                        entry = {"text": line}
                        year_match = re.search(r"\b(19|20)\d{2}\b", line)
                        if year_match:
                                entry["year"] = year_match.group(0)

                        if DEGREE_PATTERN.search(line):
                                entry["degree_match"] = DEGREE_PATTERN.search(line).group(0)

                        if not INSTITUTION_PATTERN.search(line):
                                if i + 1 < len(lines) and INSTITUTION_PATTERN.search(lines[i + 1]):
                                        entry["institution_line"] = lines[i + 1]
                                elif i > 0 and INSTITUTION_PATTERN.search(lines[i - 1]):
                                        entry["institution_line"] = lines[i - 1]

                        education.append(entry)

        deduped = []
        seen = set()
        for item in education:
                key = json.dumps(item, sort_keys=True)
                if key not in seen:
                        seen.add(key)
                        deduped.append(item)
        return deduped[:10]


def extract_certifications(lines: list[str]) -> list[str]:
        certs = []
        cert_keywords = [
                "certified", "certification", "aws certified", "azure certification",
                "pmp", "csm", "scrum master", "ccna", "security+", "itil"
        ]
        for line in lines:
                lower = line.lower()
                if any(keyword in lower for keyword in cert_keywords):
                        certs.append(line)
        return unique(certs[:20])


def parse_date_token(token: str) -> tuple[int, int] | None:
        token = token.strip().lower()
        if token in {"present", "current", "now"}:
                now = datetime.now()
                return now.year, now.month

        parts = token.split()
        if len(parts) == 1 and parts[0].isdigit():
                year = int(parts[0])
                return year, 1

        if len(parts) == 2 and parts[1].isdigit():
                month = MONTHS.get(parts[0].lower())
                year = int(parts[1])
                if month:
                        return year, month

        return None


def month_index(year: int, month: int) -> int:
        return year * 12 + (month - 1)


def extract_years_of_experience(text: str) -> float | None:
        intervals: list[tuple[int, int]] = []

        for match in DATE_RANGE_PATTERN.finditer(text):
                start_token = match.group("start")
                end_token = match.group("end")

                start = parse_date_token(start_token)
                end = parse_date_token(end_token)

                if not start or not end:
                        continue

                start_idx = month_index(*start)
                end_idx = month_index(*end)

                if end_idx < start_idx:
                        continue

                intervals.append((start_idx, end_idx))

        if not intervals:
                return None

        intervals.sort()
        merged = [intervals[0]]
        for current_start, current_end in intervals[1:]:
                last_start, last_end = merged[-1]
                if current_start <= last_end + 1:
                        merged[-1] = (last_start, max(last_end, current_end))
                else:
                        merged.append((current_start, current_end))

        total_months = sum((end - start + 1) for start, end in merged)
        return round(total_months / 12.0, 1)


def extract_summary(lines: list[str]) -> str | None:
        headings = {"summary", "professional summary", "profile", "objective"}
        for i, line in enumerate(lines[:50]):
                if line.lower() in headings:
                        summary_lines = []
                        for j in range(i + 1, min(i + 6, len(lines))):
                                candidate = lines[j].strip()
                                if len(candidate) < 3:
                                        break
                                if candidate.lower() in {
                                        "experience", "education", "skills", "projects", "certifications"
                                }:
                                        break
                                summary_lines.append(candidate)
                        if summary_lines:
                                return " ".join(summary_lines)
        return None


def process_resume(path: Path) -> dict[str, Any]:
        result: dict[str, Any] = {
                "file_name": path.name,
                "file_path": str(path.resolve()),
                "file_type": path.suffix.lower().lstrip("."),
                "name": None,
                "emails": [],
                "phones": [],
                "location": None,
                "links": {},
                "skills": [],
                "languages": [],
                "education": [],
                "certifications": [],
                "summary": None,
                "years_of_experience": None,
                "text_length": 0,
                "extraction_error": None,
        }

        try:
                text = extract_text(path)
                lines = clean_lines(text)

                result["text_length"] = len(text)
                result["name"] = extract_name(lines)
                result["emails"] = extract_emails(text)
                result["phones"] = extract_phones(text)
                result["location"] = extract_location(text, lines)

                urls = extract_urls(text)
                result["links"] = extract_links(urls)
                result["skills"] = extract_skills(text)
                result["languages"] = extract_languages(text)
                result["education"] = extract_education(lines)
                result["certifications"] = extract_certifications(lines)
                result["summary"] = extract_summary(lines)
                result["years_of_experience"] = extract_years_of_experience(text)

        except Exception as exc:
                result["extraction_error"] = str(exc)

        return result


def scan_folder(folder: Path) -> list[Path]:
        files = []
        for path in folder.rglob("*"):
                if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS and not path.name.startswith("~$"):
                        files.append(path)
        return sorted(files)


def main() -> None:
        parser = argparse.ArgumentParser(description="Extract structured data from resumes.")
        parser.add_argument("input_folder", type=str, help="Folder containing resume files")
        parser.add_argument("-o", "--output", type=str, default="extracted_resumes.json", help="Output JSON file")
        args = parser.parse_args()

        input_folder = Path(args.input_folder)
        output_file = Path(args.output)

        if not input_folder.exists() or not input_folder.is_dir():
                raise SystemExit(f"Invalid folder: {input_folder}")

        resume_files = scan_folder(input_folder)
        results = [process_resume(path) for path in resume_files]

        output_file.write_text(json.dumps(results, indent=2, ensure_ascii=False), encoding="utf-8")
        print(f"Processed {len(results)} files")
        print(f"Saved: {output_file}")


if __name__ == "__main__":
        main()